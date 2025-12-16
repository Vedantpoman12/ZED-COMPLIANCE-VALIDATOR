
import docx
import os


def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Read Paragraphs
        full_text.append("--- PARAGRAPHS ---")
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Read Tables
        full_text.append("\n--- TABLES ---")
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                full_text.append(" | ".join(row_text))
            full_text.append("-" * 20) # Separator between tables
            
        return '\n'.join(full_text)
    except Exception as e:
        return str(e)


file_path = "Pinnacle 3_Mini_Project_Report_Template_BTech.docx"
output_file = "pinnacle_content.txt"

if os.path.exists(file_path):
    print(f"Reading {file_path}...")
    content = read_docx(file_path)
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"Content written to {output_file}")
else:
    print(f"File not found: {file_path}")
